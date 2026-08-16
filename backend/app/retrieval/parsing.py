"""Document parsers: bytes to structured text."""

from __future__ import annotations

from pathlib import Path

from app.core.errors import ValidationError
from app.retrieval.protocols import DocumentParser, ParsedDocument

_ALLOWED_EXTENSIONS = frozenset({".pdf", ".docx", ".md", ".txt"})


def allowed_extension(filename: str) -> bool:
    return Path(filename).suffix.lower() in _ALLOWED_EXTENSIONS


def sniff_content_type(filename: str, data: bytes) -> str:
    """Infer a MIME type from magic bytes and extension; never trust the client."""
    ext = Path(filename).suffix.lower()
    if data.startswith(b"%PDF"):
        return "application/pdf"
    if data.startswith(b"PK\x03\x04") and ext == ".docx":
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if ext in {".md", ".txt"}:
        return "text/plain"
    if ext == ".pdf":
        return "application/pdf"
    if ext == ".docx":
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    raise ValidationError("Unsupported file type")


class PlainTextParser:
    def supports(self, content_type: str, filename: str) -> bool:
        ext = Path(filename).suffix.lower()
        return content_type.startswith("text/") or ext in {".txt", ".md"}

    def parse(self, data: bytes, *, filename: str, content_type: str) -> ParsedDocument:
        text = data.decode("utf-8", errors="replace").strip()
        if not text:
            raise ValidationError("Document contains no readable text")
        return ParsedDocument(text=text, title=Path(filename).stem)


class PdfParser:
    def supports(self, content_type: str, filename: str) -> bool:
        return content_type == "application/pdf" or Path(filename).suffix.lower() == ".pdf"

    def parse(self, data: bytes, *, filename: str, content_type: str) -> ParsedDocument:
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise ValidationError("PDF parsing requires the optional 'rag' extra (pypdf)") from exc

        from io import BytesIO

        reader = PdfReader(BytesIO(data))
        pages = [(page.extract_text() or "").strip() for page in reader.pages]
        text = "\n\n".join(page for page in pages if page)
        if not text:
            raise ValidationError("PDF contains no extractable text")
        return ParsedDocument(text=text, title=Path(filename).stem, sections=pages)


class DocxParser:
    def supports(self, content_type: str, filename: str) -> bool:
        ext = Path(filename).suffix.lower()
        return (
            content_type
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            or ext == ".docx"
        )

    def parse(self, data: bytes, *, filename: str, content_type: str) -> ParsedDocument:
        try:
            from docx import Document as DocxDocument
        except ImportError as exc:
            raise ValidationError(
                "DOCX parsing requires the optional 'rag' extra (python-docx)"
            ) from exc

        from io import BytesIO

        document = DocxDocument(BytesIO(data))
        paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
        text = "\n\n".join(paragraphs)
        if not text:
            raise ValidationError("DOCX contains no readable text")
        return ParsedDocument(text=text, title=Path(filename).stem, sections=paragraphs)


class CompositeDocumentParser:
    """Route to the first parser that accepts the file."""

    def __init__(self, parsers: list[DocumentParser] | None = None) -> None:
        self._parsers = parsers or [PlainTextParser(), PdfParser(), DocxParser()]

    def supports(self, content_type: str, filename: str) -> bool:
        return any(parser.supports(content_type, filename) for parser in self._parsers)

    def parse(self, data: bytes, *, filename: str, content_type: str) -> ParsedDocument:
        for parser in self._parsers:
            if parser.supports(content_type, filename):
                return parser.parse(data, filename=filename, content_type=content_type)
        raise ValidationError("Unsupported file type")
